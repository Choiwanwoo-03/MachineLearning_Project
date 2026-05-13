import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 설정 (A4, 여백 2.5cm) ──────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── 색상 정의 ────────────────────────────────────────────
C_TITLE   = RGBColor(0x1F, 0x49, 0x7D)   # 진남색
C_H1      = RGBColor(0x1F, 0x49, 0x7D)
C_H2      = RGBColor(0x2E, 0x74, 0xB5)
C_H3      = RGBColor(0x2E, 0x74, 0xB5)
C_TH_BG   = '1F497D'                      # 표 헤더 배경
C_TH2_BG  = '2E74B5'
C_STRIPE  = 'DEEAF1'
WHITE     = 'FFFFFF'
C_BODY    = RGBColor(0x26, 0x26, 0x26)

FONT = '맑은 고딕'


# ── 헬퍼 함수 ────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='CCCCCC', sz='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    sz)
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def para(text='', size=11, bold=False, color=None, align=None, space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.name  = FONT
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    return p

def heading1(text):
    p = para(text, size=15, bold=True, color=C_H1, space_before=18, space_after=8)
    # 하단 테두리
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    return para(text, size=12, bold=True, color=C_H2, space_before=12, space_after=4)

def heading3(text):
    return para(text, size=11, bold=True, color=C_H3, space_before=8, space_after=3)

def body(text, indent=0):
    p = para(text, size=10.5, color=C_BODY, space_before=2, space_after=4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def bullet(text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_BODY
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def make_table(headers, rows, col_widths, header_bg=C_TH_BG, stripe=True):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'

    # 헤더
    hdr = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name  = FONT
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = C_STRIPE if (stripe and ri % 2 == 0) else WHITE
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            cell.width = Cm(w)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name  = FONT
            run.font.size  = Pt(9.5)
            run.font.color.rgb = C_BODY
    return tbl


# ════════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = para('수원시 아파트 실거래가 예측 모델', size=22, bold=True,
         color=C_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
p = para('데이터 분석 결과 보고서', size=16, bold=False,
         color=C_H2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

doc.add_paragraph()
doc.add_paragraph()

info = [
    ('분석 대상', '경기도 수원시 아파트 매매 실거래'),
    ('분석 기간', '2006년 1월 ~ 2024년 12월 (19년)'),
    ('총 거래 건수', '267,319건'),
    ('독립변수', '23개 (5개 카테고리)'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{label}: ')
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = C_H2
    run2 = p.add_run(value)
    run2.font.name = FONT
    run2.font.size = Pt(11)
    run2.font.color.rgb = C_BODY

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 1. 프로젝트 진행을 위한 데이터 수집
# ════════════════════════════════════════════════════════════
heading1('1. 프로젝트 진행을 위한 데이터 수집')

heading2('1.1 프로젝트 개요')
body('본 프로젝트는 경기도 수원시 아파트 매매 실거래가를 예측하는 머신러닝 모델 개발을 목표로 한다. '
     '2006년부터 2024년까지 19년간의 실거래 데이터 267,319건을 기반으로, '
     '단지 내부, 교통 인프라, 교육·학군, 생활·환경, 거시·정책의 5개 카테고리에서 '
     '총 23개 독립변수를 수집·전처리하여 모델 학습에 활용한다.')

heading2('1.2 분석 대상 데이터 개요')

make_table(
    ['항목', '내용'],
    [
        ['분석 대상 지역', '경기도 수원시 (권선구·영통구·장안구·팔달구)'],
        ['분석 기간', '2006년 1월 ~ 2024년 12월 (19년)'],
        ['총 거래 건수', '267,319건'],
        ['고유 단지 수', '607개'],
        ['종속변수', 'price_manwon (거래금액, 만원)'],
        ['독립변수', '23개 (5개 카테고리)'],
        ['최종 데이터셋', '267,319행 × 5개 feature parquet 파일'],
    ],
    [5.5, 10.0],
)
divider()

heading2('1.3 카테고리별 데이터 수집 현황')
body('수집한 23개 변수는 5개 카테고리로 구분되며, 각 카테고리의 데이터 소스와 수집 방법은 다음과 같다.')
divider()

# 카테고리 표
heading3('가. 단지 내부 변수 (8개)')
make_table(
    ['변수명', '컬럼명', '데이터 소스', '수집 방법'],
    [
        ['전용면적',      'exclusive_area',  '국토교통부 실거래가 API',                    'RTMSDataSvcAptTradeDev excluUseAr 필드'],
        ['건축연도',      'build_year',      '국토교통부 실거래가 API',                    'RTMSDataSvcAptTradeDev buildYear 필드'],
        ['노후도',        'age',             '파생 (실거래가 API)',                        'dealYear - build_year 계산'],
        ['재건축연한더미', 'redev_dummy',     '파생 (실거래가 API)',                        'age >= 30 이진 변환'],
        ['층수',          'floor',           '국토교통부 실거래가 API',                    'RTMSDataSvcAptTradeDev floor 필드'],
        ['단지세대수',    'total_household', '경기도 공동주택 현황 (gyeonggi_apartments.csv)', '단지명 fuzzy 매핑 후 세대수 추출'],
        ['주차비율',      'parking_ratio',   '수원시 건축물대장 표제부 CSV (4개 구)',      '동별 주차대수/세대수 집계 후 단지 단위 산출'],
        ['엘리베이터유무','has_elevator',    '수원시 건축물대장 표제부 CSV (4개 구)',      '승강기수 > 0 이진 변환, 건축법 보정 적용'],
    ],
    [2.5, 3.2, 4.8, 5.0],
    header_bg=C_TH_BG,
)
divider()

heading3('나. 교통 인프라 변수 (2개)')
make_table(
    ['변수명', '컬럼명', '데이터 소스', '수집 방법'],
    [
        ['최근접 개통역 거리', 'nearest_open_dist_m', 'suwon_features.parquet (사전 계산)', '거래 시점 기준 개통역 필터링 후 Haversine 최근접 거리'],
        ['개통역 접근성 점수', 'nearest_open_score',  'suwon_features.parquet (사전 계산)', '100 / log(도보분 + 2) 로그 접근성 점수'],
    ],
    [2.8, 3.5, 4.2, 5.0],
)
divider()

heading3('다. 교육·학군 변수 (5개)')
make_table(
    ['변수명', '컬럼명', '데이터 소스', '수집 방법'],
    [
        ['초등학교 도보거리', 'elem_nearest_m',    'suwon_features.parquet (카카오 로컬 API)',   '500m 내 최근접 학교 거리, 없으면 999m 패널티'],
        ['초등학교 수(500m)', 'elem_cnt_500m',     'suwon_features.parquet (카카오 로컬 API)',   '반경 500m 이내 초등학교 수'],
        ['중학교 수(500m)',   'mid_cnt_500m',      'suwon_features.parquet (카카오 로컬 API)',   '반경 500m 이내 중학교 수'],
        ['고등학교 수(500m)', 'high_cnt_500m',     'suwon_features.parquet (카카오 로컬 API)',   '반경 500m 이내 고등학교 수'],
        ['학원 수(500m)',     'academy_cnt_500m_t','소상공인시장진흥공단 상가(상권)정보 경기 202512.csv', '수원시 학원·교습 업종 Haversine 500m 카운트'],
    ],
    [2.8, 3.5, 4.0, 5.2],
)
divider()

heading3('라. 생활·환경 변수 (4개)')
make_table(
    ['변수명', '컬럼명', '데이터 소스', '수집 방법'],
    [
        ['대형공원 거리', 'large_park_dist_m', '경기도_수원시_도시공원정보.csv', '1만㎡ 이상 공원 필터 후 Haversine 최근접 거리'],
        ['대형마트 거리', 'mart_dist_m',       '카카오 로컬 API (MT1)',         '수원 중심 15km 반경, 대형마트 필터, Haversine 최근접'],
        ['편의점 수(500m)', 'conv_cnt_500m',   '카카오 로컬 API (CS2)',         '단지별 반경 500m total_count'],
        ['종합병원 거리', 'hospital_dist_m',   '건강보험심사평가원 병원정보서비스(2026.3.).xlsx', '수원시 종합병원·상급종합 7개 Haversine 최근접'],
    ],
    [2.5, 3.2, 5.3, 4.5],
)
divider()

heading3('마. 거시·정책 변수 (4개)')
make_table(
    ['변수명', '컬럼명', '데이터 소스', '수집 방법'],
    [
        ['기준금리',      'base_rate',      '한국은행 ECOS API',      '통계코드 722Y001/0101000, 월별 조회'],
        ['주담대 금리',   'mortgage_rate',  '한국은행 ECOS API',      '통계코드 121Y006/BECBLA0302, 월별 조회'],
        ['매매가격지수',  'reb_price_idx',  '한국부동산원 R-ONE',     '수원시 아파트 매매 실거래가격지수 (2017.11=100)'],
        ['거래연도',      'deal_year',      '파생 (실거래가 API)',    'dealYear(str) → int 변환 (2006~2024)'],
    ],
    [2.5, 3.0, 4.0, 6.0],
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 2. 데이터 분석과 전처리
# ════════════════════════════════════════════════════════════
heading1('2. 데이터 분석과 전처리')

heading2('2.1 원본 데이터 현황')
body('모든 전처리 스크립트의 기반이 되는 suwon_features.parquet은 국토교통부 실거래가 API 및 카카오 로컬 API 등을 통해 사전 구축된 통합 데이터셋으로, '
     '267,319행 × 244컬럼으로 구성된다. 각 전처리 스크립트는 이 파일에서 필요한 컬럼을 추출하여 카테고리별 feature parquet을 생성한다.')

make_table(
    ['구분', '내용'],
    [
        ['베이스 파일', 'suwon_features.parquet (267,319행 × 244컬럼)'],
        ['출력 파일 수', '5개 feature parquet (카테고리별)'],
        ['총 거래 건수', '267,319건 (전처리 전후 동일)'],
        ['분석 구역', '수원시 4개 구: 영통구(37.2%) · 권선구(26.1%) · 장안구(23.8%) · 팔달구(12.9%)'],
        ['분석 기간', '2006년 ~ 2024년 (19년간)'],
    ],
    [4.0, 11.5],
)
divider()

heading2('2.2 전처리 방법론')

heading3('가. 결측치 처리 — 3단계 계층적 대체')
body('매핑 실패 또는 데이터 부재로 발생하는 결측치는 아래 3단계 순서로 처리하여 '
     '전처리 후 모든 변수의 결측치를 0건으로 만든다.')
bullet('① 1단계: 구(區) × 거래연도별 중앙값으로 대체')
bullet('② 2단계: 구(區)별 중앙값으로 대체 (1단계 후 잔여 결측)')
bullet('③ 3단계: 전체 중앙값으로 대체 (최종 안전망)')

heading3('나. 이상값 처리')
bullet('parking_ratio: 상위 1% 윈저라이징 (극단값 제한, max=12.65대/세대)')
bullet('build_year: 1900년 미만 또는 2025년 초과 → 중앙값(1999년) 대체')
bullet('age: 음수(미래 건물) → 0으로 하한 처리 (187건)')

heading3('다. 피처 엔지니어링 — 파생변수 생성')
make_table(
    ['원변수', '파생변수', '변환 방법'],
    [
        ['exclusive_area', 'log_exclusive_area, area_cat', 'log 변환; 5구간 분류(소형~대형)'],
        ['build_year',     'era_pre1990 / era_1990s / era_2000s / era_2010plus', '연대별 더미 4개'],
        ['age',            'log_age, redev_dummy',    'log(age+1); 30년 이상=1 재건축 더미'],
        ['floor',          'log_floor, is_ground_floor', 'log 변환; 1층=1 더미'],
        ['total_household','log_total_household',    'log 변환'],
        ['parking_ratio',  'has_parking',            '주차대수 > 0 이진 더미'],
        ['nearest_open_dist_m','nearest_open_score', '100/log(도보분+2) 로그 접근성 점수'],
        ['elem_nearest_m', 'log_elem_nearest_m, elem_access_score, elem_walkable_500m', 'log; 접근성 점수; 500m 이내 더미'],
        ['elem_cnt_500m 외 2개', '학교 밀도 더미(has_elem_500m 등)', '1개 이상 존재 이진 더미'],
    ],
    [3.5, 5.0, 7.0],
)
divider()

heading2('2.3 카테고리별 전처리 상세')

heading3('가. 단지 내부 변수')
body('국토교통부 실거래가 API에서 제공되는 dealAmount(콤마 포함 문자열)를 정수 변환하여 '
     '타겟 변수 price_manwon을 생성한다. 엘리베이터 유무는 건축물대장 표제부와 fuzzy 매핑 후 '
     '건축법 제64조(6층 이상 설치 의무)를 적용하여 84,445건을 자동 보정하였다.')
bullet('단지명 fuzzy 매핑 결과: total_household 매핑 성공률 54.7%, parking_ratio 54.4%')
bullet('매핑 실패 단지는 구×연도별 중앙값으로 imputation 처리')

heading3('나. 교통 인프라 변수')
body('거래 시점 기준 실제 개통된 역만을 대상으로 최근접 거리를 산출하여 '
     'temporal leakage를 방지한다. 2014년 수인선 수원~오이도 구간 개통 이후 '
     '평균 거리가 2,195m에서 1,484m로 감소하는 시계열 패턴이 확인된다.')

heading3('다. 교육·학군 변수')
body('학교 관련 4개 변수(elem_nearest_m, elem/mid/high_cnt_500m)는 suwon_features.parquet에 '
     '카카오 로컬 API 기반으로 사전 수집된 값을 추출한다. '
     'elem_nearest_m의 NaN(500m 내 학교 없음)은 999m 패널티로 처리하며, '
     'elem_walkable_500m 더미를 병행 생성한다. '
     'academy_cnt_500m_t는 2025.12 정적 스냅샷 기준이므로 시계열 누수(temporal leakage) 주의가 필요하다.')

heading3('라. 생활·환경 변수')
body('카카오 로컬 API 결과는 mart_locations.csv, conv_cnt_lookup.csv에 캐시하여 재수집을 방지한다. '
     '카카오 API 키는 환경변수(KAKAO_API_KEY)로 관리하며, 캐시 파일이 없을 경우에만 API를 호출한다. '
     '4개 변수 모두 정적 스냅샷(현재 시점 기준)이다.')

heading3('마. 거시·정책 변수')
body('한국은행 ECOS Open API(인증키 필요)를 통해 월별 기준금리와 주담대 금리를 수집하며, '
     '한국부동산원 R-ONE에서 수원시 아파트 매매 실거래가격지수를 취득한다. '
     '3개 시계열 변수는 거래 월(_ym) 기준 LEFT JOIN으로 병합하여 결측 0건을 달성한다.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 3. 데이터 분석 결과 보고서
# ════════════════════════════════════════════════════════════
heading1('3. 프로젝트 목표 달성을 위한 데이터 분석 결과')

heading2('3.1 최종 데이터셋 현황')

make_table(
    ['파일명', '변수 수', '행 수', '결측치'],
    [
        ['complex_inner_features.parquet', '28개 (ID 9 + 피처 19)', '267,319', '0건'],
        ['traffic_features.parquet',       '11개 (ID 9 + 피처 2)',  '267,319', '0건'],
        ['edu_features.parquet',           '25개 (ID 9 + 피처 16)', '267,319', '0건'],
        ['env_features.parquet',           '13개 (ID 9 + 피처 4)',  '267,319', '0건'],
        ['macro_features.parquet',         '13개 (ID 9 + 피처 4)',  '267,319', '0건'],
    ],
    [5.5, 4.5, 2.5, 2.0],
)
divider()

heading2('3.2 종속변수 (price_manwon) 분석')
body('종속변수인 아파트 매매 거래금액(만원)의 기초 통계량은 다음과 같다.')

make_table(
    ['통계량', '값'],
    [
        ['관측 수', '267,319건'],
        ['평균',    '31,011만원 (약 3.1억)'],
        ['표준편차','19,595만원'],
        ['최솟값',  '2,950만원'],
        ['25 백분위수', '18,400만원'],
        ['중앙값',  '25,900만원 (약 2.6억)'],
        ['75 백분위수', '37,900만원'],
        ['최댓값',  '338,000만원 (약 33.8억)'],
    ],
    [5.0, 5.0],
)
divider()

heading3('연도별 거래 건수')
make_table(
    ['연도', '거래 건수', '연도', '거래 건수', '연도', '거래 건수'],
    [
        ['2006', '23,222', '2013', '13,967', '2020', '21,984'],
        ['2007',  '9,276', '2014', '18,733', '2021', '13,691'],
        ['2008', '10,746', '2015', '17,897', '2022',  '3,526'],
        ['2009', '14,493', '2016', '16,332', '2023',  '9,456'],
        ['2010', '10,551', '2017', '15,005', '2024', '11,795'],
        ['2011', '14,163', '2018', '16,828', '합계', '267,319'],
        ['2012',  '8,942', '2019', '16,712', '-',        '-'],
    ],
    [2.0, 2.5, 2.0, 2.5, 2.0, 2.5],
)
divider()

heading2('3.3 주요 독립변수 기초 통계')

heading3('가. 단지 내부 변수')
make_table(
    ['변수명', '평균', '중앙값', '최솟값', '최댓값', '비고'],
    [
        ['exclusive_area (㎡)', '73.59', '60.00', '10.93', '229.82', '중소형(33~60㎡) 49.0%'],
        ['build_year',          '2000년', '1999년', '1978년', '2024년', '1990년대 건물 43.6%'],
        ['age (년)',             '14.6', '14.0', '0', '46', '재건축 대상(30년+): 4.6%'],
        ['floor',               '9.4', '9.0', '1', '49', '1층 거래: 15,408건'],
        ['total_household (세대)', '1,021', '797', '30', '5,282', '매핑 성공률 54.7%'],
        ['parking_ratio (대/세대)', '0.941', '0.000', '0.000', '12.649', 'has_parking=1: 28.1%'],
        ['has_elevator',        '-', '-', '0', '1', '엘리베이터 있음: 81.1%'],
    ],
    [3.8, 1.8, 1.8, 1.8, 1.8, 4.5],
)
divider()

heading3('나. 교통 인프라 변수')
make_table(
    ['변수명', '평균', '중앙값', '최솟값', '최댓값'],
    [
        ['nearest_open_dist_m (m)', '1,726', '1,445', '37', '5,942'],
        ['nearest_open_score (점)', '30.7', '29.4', '-', '-'],
    ],
    [4.5, 2.5, 2.5, 2.5, 2.5],
)
divider()

heading3('다. 교육·학군 변수')
make_table(
    ['변수명', '평균', '최솟값', '최댓값', '비고'],
    [
        ['elem_nearest_m (m)',    '338',  '43',  '999',  '500m 내 학교 없음(999): 14.8%'],
        ['elem_cnt_500m (개)',    '1.39', '0',   '5',    '500m 내 초등학교 있음: 82.0%'],
        ['mid_cnt_500m (개)',     '0.85', '0',   '3',    '500m 내 중학교 있음: 64.2%'],
        ['high_cnt_500m (개)',    '0.50', '0',   '3',    '500m 내 고등학교 있음: 42.4%'],
        ['academy_cnt_500m_t (개)', '63.4', '0', '263', '2025.12 정적 스냅샷'],
    ],
    [4.0, 1.8, 1.8, 1.8, 6.1],
)
divider()

heading3('라. 생활·환경 변수')
make_table(
    ['변수명', '평균', '최솟값', '최댓값', '비고'],
    [
        ['large_park_dist_m (m)', '504',  '3',   '-',  '1만㎡ 이상 공원 기준'],
        ['mart_dist_m (m)',       '1,239', '-',  '-',  '대형마트 기준'],
        ['conv_cnt_500m (개)',    '12.6',  '0',  '43', '500m 반경 편의점 수'],
        ['hospital_dist_m (m)',   '1,854', '-',  '-',  '종합병원·상급종합 기준'],
    ],
    [4.2, 1.8, 1.8, 1.8, 5.9],
)
divider()

heading3('마. 거시·정책 변수')
make_table(
    ['변수명', '평균', '최솟값', '최댓값', '비고'],
    [
        ['base_rate (%)',     '2.34', '0.50', '5.25', '2021년 최저(0.50%), 2023년 최고(3.50%)'],
        ['mortgage_rate (%)', '4.01', '2.39', '7.58', '2008년 금융위기 시 최고'],
        ['reb_price_idx',    '74.2', '45.7', '109.8', '기준: 2017.11=100'],
        ['deal_year',        '-',    '2006',  '2024',  'dealYear(str) → int 변환'],
    ],
    [3.5, 1.5, 1.8, 1.8, 6.9],
)

divider()
heading2('3.4 데이터 품질 요약')
body('전처리 완료 후 5개 feature parquet 파일 모두 결측치 0건을 달성하였으며, '
     '행수는 267,319건으로 원본과 동일하게 유지된다. '
     '주요 품질 보정 내역은 아래와 같다.')

make_table(
    ['항목', '내용', '처리 건수'],
    [
        ['결측치 처리', '3단계 계층적 중앙값 대체 (구×연도→구→전체)', '전 변수 0건 달성'],
        ['이상값 보정', 'parking_ratio 상위 1% 윈저라이징', '-'],
        ['건축법 보정', '6층 이상 has_elevator=0 → 1 강제 보정', '84,445건'],
        ['음수 age 처리', 'clip(lower=0) 하한 처리', '187건'],
        ['이상 건축연도', '범위 외 값 → 중앙값(1999년) 대체', '0건 (해당 없음)'],
        ['area_cat dtype', 'category → str 변환 (parquet 호환)', '전 행'],
    ],
    [3.5, 8.0, 4.0],
)

# ── 저장 ──────────────────────────────────────────────────
out = r'C:\Users\최완우\OneDrive\Desktop\기계학습 기말 프로젝트_최한결\수원시_아파트_데이터분석_결과보고서.docx'
doc.save(out)
print(f'저장 완료: {out}')
